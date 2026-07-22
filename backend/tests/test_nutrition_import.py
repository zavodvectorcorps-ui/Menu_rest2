"""Tests for nutrition import from .docx (parser + matcher + endpoint edge cases)."""
import io
import os
import pytest
import requests
from docx import Document

BASE_URL = os.environ.get("REACT_APP_BACKEND_URL", "https://restaurant-hub-275.preview.emergentagent.com").rstrip("/")
API = f"{BASE_URL}/api"

SAMPLE_DOCX = "/tmp/bzu.docx"


# ---------- fixtures ----------
@pytest.fixture(scope="session")
def admin_token():
    r = requests.post(f"{API}/auth/login", json={"username": "admin", "password": "220066"}, timeout=15)
    assert r.status_code == 200, r.text
    return r.json()["access_token"]


@pytest.fixture(scope="session")
def auth_headers(admin_token):
    return {"Authorization": f"Bearer {admin_token}"}


@pytest.fixture(scope="session")
def restaurant_id(auth_headers):
    r = requests.get(f"{API}/restaurants", headers=auth_headers, timeout=15)
    assert r.status_code == 200
    rests = r.json()
    assert len(rests) > 0
    # Prefer non-demo
    for x in rests:
        if x.get("slug") != "demo":
            return x["id"]
    return rests[0]["id"]


@pytest.fixture(scope="session")
def category_id(auth_headers, restaurant_id):
    r = requests.get(f"{API}/restaurants/{restaurant_id}/categories", headers=auth_headers, timeout=15)
    assert r.status_code == 200
    cats = r.json()
    if cats:
        return cats[0]["id"]
    r = requests.post(f"{API}/restaurants/{restaurant_id}/categories", headers=auth_headers,
                      json={"name": "TEST_Cat"}, timeout=15)
    return r.json()["id"]


# ---------- parser unit tests ----------
def test_parser_per_dish_table():
    """Formatting: one table per dish (row0=name, last row=values)."""
    from services.nutrition_import import parse_docx_nutrition

    doc = Document()
    t = doc.add_table(rows=4, cols=5)
    t.rows[0].cells[0].text = "Сырники"
    t.rows[1].cells[0].text = "Количество в 100 граммах"
    for i, h in enumerate(["Белки", "Жиры", "Углеводы", "Ккал", "кДж"]):
        t.rows[2].cells[i].text = h
    for i, v in enumerate(["11,15", "16,37", "28,74", "306", "1285"]):
        t.rows[3].cells[i].text = v

    buf = io.BytesIO()
    doc.save(buf)
    recs = parse_docx_nutrition(buf.getvalue())
    assert len(recs) == 1
    r = recs[0]
    assert r["name"] == "Сырники"
    assert r["protein"] == 11.15
    assert r["fat"] == 16.37
    assert r["carbs"] == 28.74
    assert r["kcal"] == 306
    assert r["kj"] == 1285


def test_parser_flat_table():
    """Formatting: one flat table with header + N rows."""
    from services.nutrition_import import parse_docx_nutrition

    doc = Document()
    t = doc.add_table(rows=7, cols=6)
    hdr = ["Название", "Белки", "Жиры", "Углеводы", "Ккал", "кДж"]
    for i, h in enumerate(hdr):
        t.rows[0].cells[i].text = h
    rows = [
        ("Борщ", "3.5", "5.2", "10.1", "95", "397"),
        ("Салат Цезарь", "6.0", "12.0", "8.0", "160", "670"),
        ("Пельмени", "10.0", "9.0", "22.0", "220", "920"),
        ("Гранола", "9.1", "14.2", "60.0", "410", "1717"),
        ("Шакшука", "7.8", "8.1", "5.4", "125", "523"),
        ("Пустая", "", "", "", "", ""),
    ]
    for ri, row in enumerate(rows, start=1):
        for ci, v in enumerate(row):
            t.rows[ri].cells[ci].text = v
    buf = io.BytesIO()
    doc.save(buf)
    recs = parse_docx_nutrition(buf.getvalue())
    names = [r["name"] for r in recs]
    assert "Борщ" in names
    assert "Салат Цезарь" in names
    assert "Гранола" in names
    assert "Пустая" not in names  # rows without numbers skipped
    assert len(recs) == 5


def test_parser_real_file_if_present():
    if not os.path.exists(SAMPLE_DOCX):
        pytest.skip("no sample file")
    from services.nutrition_import import parse_docx_nutrition
    with open(SAMPLE_DOCX, "rb") as f:
        recs = parse_docx_nutrition(f.read())
    # user says 67 records
    assert len(recs) >= 30, f"only {len(recs)} records parsed"


# ---------- matcher unit tests ----------
def test_matcher_exact_and_ambiguous():
    from services.nutrition_import import match_records_to_items
    items = [
        {"id": "1", "name": "Борщ"},
        {"id": "2", "name": "Салат Цезарь с курицей"},
        {"id": "3", "name": "Салат Цезарь с креветкой"},
        {"id": "4", "name": "Пельмени домашние"},
    ]
    records = [
        {"name": "Борщ", "protein": 3.5, "fat": 5.2, "carbs": 10.1, "kcal": 95, "kj": 397},
        {"name": "Салат Цезарь", "protein": 6, "fat": 12, "carbs": 8, "kcal": 160, "kj": 670},
        {"name": "Совершенно другое блюдо xyz", "protein": 1, "fat": 1, "carbs": 1, "kcal": 1, "kj": 1},
    ]
    res = match_records_to_items(records, items)
    matched_sources = [m["source"] for m in res["matched"]]
    assert "Борщ" in matched_sources
    # Салат Цезарь — two close candidates should be ambiguous
    amb_sources = [a["source"] for a in res["ambiguous"]]
    unmatched_sources = [u["source"] for u in res["unmatched"]]
    assert ("Салат Цезарь" in amb_sources) or ("Салат Цезарь" in matched_sources)
    assert "Совершенно другое блюдо xyz" in unmatched_sources


# ---------- endpoint edge cases ----------
def _make_flat_docx(dish_names_with_vals):
    doc = Document()
    t = doc.add_table(rows=1 + len(dish_names_with_vals), cols=6)
    for i, h in enumerate(["Название", "Белки", "Жиры", "Углеводы", "Ккал", "кДж"]):
        t.rows[0].cells[i].text = h
    for ri, (name, vals) in enumerate(dish_names_with_vals, start=1):
        t.rows[ri].cells[0].text = name
        for ci, v in enumerate(vals):
            t.rows[ri].cells[ci + 1].text = str(v)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_endpoint_rejects_non_docx(auth_headers, restaurant_id):
    files = {"file": ("test.txt", b"hello", "text/plain")}
    r = requests.post(
        f"{API}/restaurants/{restaurant_id}/menu-items/nutrition-import",
        headers=auth_headers, files=files, params={"dry_run": "true"}, timeout=30,
    )
    assert r.status_code == 400, r.text


def test_endpoint_rejects_empty_file(auth_headers, restaurant_id):
    files = {"file": ("empty.docx", b"", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    r = requests.post(
        f"{API}/restaurants/{restaurant_id}/menu-items/nutrition-import",
        headers=auth_headers, files=files, params={"dry_run": "true"}, timeout=30,
    )
    assert r.status_code == 400


def test_endpoint_rejects_docx_without_tables(auth_headers, restaurant_id):
    doc = Document()
    doc.add_paragraph("No tables here")
    buf = io.BytesIO()
    doc.save(buf)
    files = {"file": ("notable.docx", buf.getvalue(),
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    r = requests.post(
        f"{API}/restaurants/{restaurant_id}/menu-items/nutrition-import",
        headers=auth_headers, files=files, params={"dry_run": "true"}, timeout=30,
    )
    assert r.status_code == 400
    assert "не найдено" in r.text.lower() or "not found" in r.text.lower() or "ни одной" in r.text.lower()


def test_endpoint_dry_run_preview_shape(auth_headers, restaurant_id):
    # Get some real dish names to guarantee at least one matched
    r = requests.get(f"{API}/restaurants/{restaurant_id}/menu-items", headers=auth_headers, timeout=15)
    assert r.status_code == 200
    items = [x for x in r.json() if not x.get("is_banner")]
    assert len(items) >= 2
    known = items[0]["name"]
    # Use >= 5 data rows so flat-table branch triggers (per-dish branch = 3..5 rows total)
    doc_bytes = _make_flat_docx([
        (known, [1.1, 2.2, 3.3, 100, 418]),
        ("ZZZ Абсолютно несуществующее блюдо 999", [1, 1, 1, 1, 1]),
        ("Filler1", [1, 1, 1, 1, 1]),
        ("Filler2", [1, 1, 1, 1, 1]),
        ("Filler3", [1, 1, 1, 1, 1]),
    ])
    files = {"file": ("test.docx", doc_bytes,
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    r = requests.post(
        f"{API}/restaurants/{restaurant_id}/menu-items/nutrition-import",
        headers=auth_headers, files=files, params={"dry_run": "true"}, timeout=30,
    )
    assert r.status_code == 200, r.text
    data = r.json()
    assert data["records_total"] == 5
    assert set(data.keys()) >= {"matched", "ambiguous", "unmatched", "records_total", "applied"}
    assert data["applied"] == 0
    all_sources = (
        [m["source"] for m in data["matched"]]
        + [a["source"] for a in data["ambiguous"]]
        + [u["source"] for u in data["unmatched"]]
    )
    assert known in all_sources
    assert any("ZZZ" in s for s in [u["source"] for u in data["unmatched"]])


def test_endpoint_apply_with_apply_ids_persists(auth_headers, restaurant_id, category_id):
    # Create a fresh dish to safely overwrite
    payload = {
        "name": "TEST_NutritionImportDish",
        "price": 100,
        "category_id": category_id,
    }
    r = requests.post(f"{API}/restaurants/{restaurant_id}/menu-items", headers=auth_headers, json=payload, timeout=15)
    assert r.status_code in (200, 201), r.text
    created = r.json()
    item_id = created["id"]

    try:
        doc_bytes = _make_flat_docx([
            ("TEST_NutritionImportDish", [7.7, 8.8, 9.9, 250, 1046]),
        ])
        files = {"file": ("apply.docx", doc_bytes,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        r = requests.post(
            f"{API}/restaurants/{restaurant_id}/menu-items/nutrition-import",
            headers=auth_headers, files=files,
            params={"dry_run": "false", "apply_ids": item_id},
            timeout=30,
        )
        assert r.status_code == 200, r.text
        data = r.json()
        assert data["applied"] >= 1

        # Verify persistence
        r = requests.get(f"{API}/restaurants/{restaurant_id}/menu-items", headers=auth_headers, timeout=15)
        got = next((x for x in r.json() if x["id"] == item_id), None)
        assert got is not None
        assert abs(got.get("nutrition_protein", 0) - 7.7) < 1e-6
        assert abs(got.get("nutrition_fat", 0) - 8.8) < 1e-6
        assert abs(got.get("nutrition_carbs", 0) - 9.9) < 1e-6
        assert abs(got.get("nutrition_kcal", 0) - 250) < 1e-6
        assert abs(got.get("nutrition_kj", 0) - 1046) < 1e-6
    finally:
        requests.delete(f"{API}/restaurants/{restaurant_id}/menu-items/{item_id}", headers=auth_headers, timeout=15)


def test_endpoint_apply_ids_filter_excludes_others(auth_headers, restaurant_id, category_id):
    """When apply_ids given, only that item is updated."""
    # Create two dishes
    ids = []
    for name in ["TEST_NIA", "TEST_NIB"]:
        r = requests.post(f"{API}/restaurants/{restaurant_id}/menu-items", headers=auth_headers,
                          json={"name": name, "price": 100, "category_id": category_id}, timeout=15)
        assert r.status_code in (200, 201)
        ids.append(r.json()["id"])
    try:
        doc_bytes = _make_flat_docx([
            ("TEST_NIA", [1.0, 2.0, 3.0, 50, 209]),
            ("TEST_NIB", [4.0, 5.0, 6.0, 80, 335]),
            ("Filler1", [1, 1, 1, 1, 1]),
            ("Filler2", [1, 1, 1, 1, 1]),
            ("Filler3", [1, 1, 1, 1, 1]),
        ])
        files = {"file": ("apply.docx", doc_bytes,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        r = requests.post(
            f"{API}/restaurants/{restaurant_id}/menu-items/nutrition-import",
            headers=auth_headers, files=files,
            params={"dry_run": "false", "apply_ids": ids[0]},  # only first
            timeout=30,
        )
        assert r.status_code == 200
        assert r.json()["applied"] == 1

        r = requests.get(f"{API}/restaurants/{restaurant_id}/menu-items", headers=auth_headers, timeout=15)
        by_id = {x["id"]: x for x in r.json()}
        assert by_id[ids[0]].get("nutrition_protein") == 1.0
        assert by_id[ids[1]].get("nutrition_protein") in (None, 0, 0.0)
    finally:
        for iid in ids:
            requests.delete(f"{API}/restaurants/{restaurant_id}/menu-items/{iid}", headers=auth_headers, timeout=15)
